#!/usr/bin/env python3
"""
Markdown to DOCX Converter with Mermaid Diagram Support

Converts markdown files to properly formatted DOCX documents, including:
- Headings, lists, tables, code blocks, links, images
- Mermaid diagrams (converted to images)

Usage:
    python md_to_docx.py input.md [output.docx]
"""

import argparse
import base64
import io
import os
import re
import subprocess
import sys
import tempfile
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx")
    sys.exit(1)

try:
    import markdown
    from markdown.extensions import tables, fenced_code, codehilite
except ImportError:
    print("Error: markdown is required. Install with: pip install markdown")
    sys.exit(1)


class MarkdownToDocxConverter:
    def __init__(self, input_file, output_file=None):
        self.input_file = Path(input_file)
        if output_file:
            self.output_file = Path(output_file)
        else:
            self.output_file = self.input_file.with_suffix('.docx')

        self.doc = Document()
        self.temp_files = []

    def cleanup(self):
        """Remove temporary files"""
        for temp_file in self.temp_files:
            try:
                if os.path.exists(temp_file):
                    os.unlink(temp_file)
            except Exception as e:
                print(f"Warning: Could not delete temp file {temp_file}: {e}")

    def extract_and_convert_mermaid(self, md_content):
        """Extract mermaid code blocks and convert them to images"""
        # Match opening fence (3+ backticks) with 'mermaid' info string,
        # then content, then closing fence (3+ backticks on its own line).
        mermaid_pattern = r'`{3,}mermaid\r?\n(.*?)\r?\n`{3,}'

        def replace_mermaid(match):
            mermaid_code = match.group(1)
            image_path = self.render_mermaid(mermaid_code)
            if image_path:
                # Replace with markdown image syntax
                return f'![Mermaid Diagram]({image_path})'
            return match.group(0)  # Return original if conversion failed

        return re.sub(mermaid_pattern, replace_mermaid, md_content, flags=re.DOTALL)

    def render_mermaid(self, mermaid_code):
        """Render mermaid diagram to PNG image"""
        # Try using mermaid-cli (mmdc)
        if self._check_command('mmdc'):
            return self._render_with_mmdc(mermaid_code)

        # Try mermaid.ink online renderer (no local dependencies)
        result = self._render_with_mermaid_ink(mermaid_code)
        if result:
            return result

        print("Warning: Mermaid rendering not available (mermaid.ink unreachable and mmdc not installed)")
        return None

    def _check_command(self, cmd):
        """Check if a command is available"""
        try:
            subprocess.run([cmd, '--version'], capture_output=True, timeout=5)
            return True
        except (subprocess.SubprocessError, FileNotFoundError):
            return False

    def _render_with_mmdc(self, mermaid_code):
        """Render using mermaid-cli (mmdc)"""
        with tempfile.NamedTemporaryFile(mode='w', suffix='.mmd', delete=False) as f:
            f.write(mermaid_code)
            input_file = f.name

        output_file = input_file.replace('.mmd', '.png')
        self.temp_files.extend([input_file, output_file])

        try:
            subprocess.run(['mmdc', '-i', input_file, '-o', output_file, '-b', 'white'],
                         check=True, capture_output=True, timeout=30)
            return output_file
        except subprocess.SubprocessError as e:
            print(f"Warning: Failed to render mermaid diagram: {e}")
            return None

    def _render_with_mermaid_ink(self, mermaid_code):
        """Render using the mermaid.ink online service (no local deps needed)."""
        import urllib.request
        import urllib.error

        try:
            # mermaid.ink accepts base64-encoded diagram definitions
            encoded = base64.urlsafe_b64encode(mermaid_code.encode('utf-8')).decode('ascii')
            url = f'https://mermaid.ink/img/{encoded}?type=png&bgColor=!FFFFFF'

            output_file = tempfile.mktemp(suffix='.png')
            self.temp_files.append(output_file)

            req = urllib.request.Request(url, headers={'User-Agent': 'md-to-docx/1.0'})
            with urllib.request.urlopen(req, timeout=30) as resp:
                data = resp.read()

            with open(output_file, 'wb') as f:
                f.write(data)

            # Sanity-check: a valid PNG is at least a few hundred bytes
            if os.path.getsize(output_file) > 100:
                return output_file

            print("Warning: mermaid.ink returned a suspiciously small image")
            return None
        except (urllib.error.URLError, urllib.error.HTTPError, OSError) as e:
            print(f"Warning: mermaid.ink rendering failed: {e}")
            return None

    def add_heading(self, text, level):
        """Add a heading to the document"""
        heading = self.doc.add_heading(text, level=level)
        return heading

    def add_paragraph(self, text, style=None):
        """Add a paragraph with inline formatting"""
        p = self.doc.add_paragraph(style=style)

        # Parse inline markdown: **bold**, *italic*, `code`, [links](url)
        parts = self._parse_inline_markdown(text)

        for part_type, content, url in parts:
            if part_type == 'bold':
                p.add_run(content).bold = True
            elif part_type == 'italic':
                p.add_run(content).italic = True
            elif part_type == 'code':
                run = p.add_run(content)
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
            elif part_type == 'link':
                self._add_hyperlink(p, url, content)
            elif part_type == 'color_icon':
                run = p.add_run(content)
                run.font.color.rgb = RGBColor.from_string(url)  # url holds hex color
                run.font.size = Pt(12)
            else:
                p.add_run(content)

        return p

    # Mapping of color emoji to hex RGB for foreground-colored ● symbols
    _COLOR_EMOJI_TO_RGB = {
        '\U0001f7e2': '4CAF50',   # 🟢 Green
        '\U0001f534': 'F44336',   # 🔴 Red
        '\U0001f535': '3F51B5',   # 🔵 Blue / Indigo
        '\U0001f7e0': 'FF9800',   # 🟠 Orange
        '\U0001f7e1': 'FFEB3B',   # 🟡 Yellow
        '\U0001f7e6': '0097A7',   # 🟦 Teal
        '\u2b1c':     '9E9E9E',   # ⬜ Gray
        '\U0001f7e8': 'F9A825',   # 🟨 Yellow (box)
        '\U0001f7e5': 'FF5722',   # 🟥 Red-orange
        '\U0001f7ea': '9C27B0',   # 🟪 Purple
        '\U0001f7e9': '4CAF50',   # 🟩 Green (box)
    }

    # Build a single regex that matches any known color emoji
    _COLOR_EMOJI_RE = re.compile('(' + '|'.join(re.escape(e) for e in _COLOR_EMOJI_TO_RGB) + ')')

    def _parse_inline_markdown(self, text):
        """Parse inline markdown formatting"""
        parts = []
        patterns = [
            (r'\*\*(.+?)\*\*', 'bold'),
            (r'\*(.+?)\*', 'italic'),
            (r'`(.+?)`', 'code'),
            (r'\[(.+?)\]\((.+?)\)', 'link'),
        ]

        pos = 0
        while pos < len(text):
            matched = False

            # Check for color emoji first
            emoji_match = self._COLOR_EMOJI_RE.match(text, pos)
            if emoji_match:
                emoji = emoji_match.group(1)
                hex_color = self._COLOR_EMOJI_TO_RGB[emoji]
                parts.append(('color_icon', '\u25CF', hex_color))  # ● with color
                pos = emoji_match.end()
                continue

            for pattern, fmt_type in patterns:
                match = re.match(pattern, text[pos:])
                if match:
                    if fmt_type == 'link':
                        parts.append((fmt_type, match.group(1), match.group(2)))
                    else:
                        parts.append((fmt_type, match.group(1), None))
                    pos += match.end()
                    matched = True
                    break

            if not matched:
                # Find next special character or emoji
                next_pos = len(text)
                for pattern, _ in patterns:
                    search = re.search(pattern, text[pos:])
                    if search and search.start() > 0:
                        next_pos = min(next_pos, pos + search.start())
                emoji_search = self._COLOR_EMOJI_RE.search(text, pos)
                if emoji_search and emoji_search.start() > pos:
                    next_pos = min(next_pos, emoji_search.start())

                if next_pos > pos:
                    parts.append(('text', text[pos:next_pos], None))
                    pos = next_pos
                else:
                    parts.append(('text', text[pos], None))
                    pos += 1

        return parts

    def _add_hyperlink(self, paragraph, url, text):
        """Add a hyperlink to a paragraph"""
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)

        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)

        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')

        # Style the hyperlink
        c = OxmlElement('w:color')
        c.set(qn('w:val'), '0563C1')
        rPr.append(c)
        u = OxmlElement('w:u')
        u.set(qn('w:val'), 'single')
        rPr.append(u)

        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)

        paragraph._p.append(hyperlink)

    def add_code_block(self, code, language=''):
        """Add a code block"""
        p = self.doc.add_paragraph(style='Normal')
        run = p.add_run(code)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

        # Add background shading
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'F5F5F5')
        p._p.get_or_add_pPr().append(shading_elm)

        return p

    def add_table(self, table_lines):
        """Add a markdown table to the document as a proper Word table.

        Expects a list of pipe-delimited lines:
            | Header1 | Header2 |
            |---------|---------|
            | Cell1   | Cell2   |
        """
        # Parse rows, stripping outer pipes
        rows = []
        separator_idx = None
        for idx, line in enumerate(table_lines):
            stripped = line.strip()
            if not stripped:
                continue
            # Detect separator row (all dashes/colons/pipes/spaces)
            if re.match(r'^[\|\s:\-]+$', stripped):
                separator_idx = idx
                continue
            cells = [c.strip() for c in stripped.strip('|').split('|')]
            rows.append(cells)

        if not rows:
            return

        # Determine column count from the widest row
        num_cols = max(len(r) for r in rows)

        # Normalise every row to the same column count
        for r in rows:
            while len(r) < num_cols:
                r.append('')

        # Create the Word table
        table = self.doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'

        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_text in enumerate(row_data):
                cell = table.rows[row_idx].cells[col_idx]
                # Clear the default empty paragraph
                cell.text = ''
                p = cell.paragraphs[0]
                # Use inline markdown parser so bold/italic/code/links work
                parts = self._parse_inline_markdown(cell_text)
                for part_type, content, url in parts:
                    if part_type == 'bold':
                        run = p.add_run(content)
                        run.bold = True
                        run.font.size = Pt(10)
                    elif part_type == 'italic':
                        run = p.add_run(content)
                        run.italic = True
                        run.font.size = Pt(10)
                    elif part_type == 'code':
                        run = p.add_run(content)
                        run.font.name = 'Courier New'
                        run.font.size = Pt(9)
                    elif part_type == 'link':
                        self._add_hyperlink(p, url, content)
                    elif part_type == 'color_icon':
                        run = p.add_run(content)
                        run.font.color.rgb = RGBColor.from_string(url)
                        run.font.size = Pt(12)
                    else:
                        run = p.add_run(content)
                        run.font.size = Pt(10)

                # Bold the header row
                if row_idx == 0 and separator_idx is not None:
                    for run in p.runs:
                        run.bold = True

    def _create_numbering(self):
        """Create a new numbered list definition and return its numId.

        Each call produces a fresh abstract + concrete numbering definition
        that starts at 1, so separate list groups never share counters.
        """
        numbering_part = self.doc.part.numbering_part
        numbering_elem = numbering_part._element

        # Find the highest existing abstractNumId
        max_abstract = -1
        for an in numbering_elem.findall(qn('w:abstractNum')):
            aid = int(an.get(qn('w:abstractNumId')))
            if aid > max_abstract:
                max_abstract = aid
        new_abstract_id = max_abstract + 1

        # Find highest existing numId
        max_num = 0
        for n in numbering_elem.findall(qn('w:num')):
            nid = int(n.get(qn('w:numId')))
            if nid > max_num:
                max_num = nid
        new_num_id = max_num + 1

        # Create a new abstractNum with a simple decimal list at level 0
        abstract_num = OxmlElement('w:abstractNum')
        abstract_num.set(qn('w:abstractNumId'), str(new_abstract_id))
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), '0')
        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)
        fmt = OxmlElement('w:numFmt')
        fmt.set(qn('w:val'), 'decimal')
        lvl.append(fmt)
        lvl_text = OxmlElement('w:lvlText')
        lvl_text.set(qn('w:val'), '%1.')
        lvl.append(lvl_text)
        lvl_jc = OxmlElement('w:lvlJc')
        lvl_jc.set(qn('w:val'), 'left')
        lvl.append(lvl_jc)
        abstract_num.append(lvl)
        numbering_elem.insert(0, abstract_num)

        # Create a new num referencing the abstract
        num = OxmlElement('w:num')
        num.set(qn('w:numId'), str(new_num_id))
        abstract_ref = OxmlElement('w:abstractNumId')
        abstract_ref.set(qn('w:val'), str(new_abstract_id))
        num.append(abstract_ref)
        numbering_elem.append(num)

        return new_num_id

    def _apply_numbering(self, paragraph, num_id):
        """Point a paragraph at a specific numbering definition."""
        pPr = paragraph._p.get_or_add_pPr()
        numPr = OxmlElement('w:numPr')
        ilvl = OxmlElement('w:ilvl')
        ilvl.set(qn('w:val'), '0')
        numPr.append(ilvl)
        numId_elem = OxmlElement('w:numId')
        numId_elem.set(qn('w:val'), str(num_id))
        numPr.append(numId_elem)
        pPr.append(numPr)

    def add_horizontal_rule(self):
        """Add a proper horizontal rule (bottom border on an empty paragraph)."""
        p = self.doc.add_paragraph()
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        pBdr.append(bottom)
        pPr.append(pBdr)

    def _apply_blockquote_style(self, element):
        """Apply blockquote styling (left indent + left border) to a paragraph or table.

        Works on both Paragraph objects and Table objects.
        """
        from docx.table import Table as DocxTable

        if isinstance(element, DocxTable):
            # For tables, only indent — don't add left border to cell paragraphs
            # (the table grid provides its own borders)
            tPr = element._tbl.tblPr
            if tPr is None:
                tPr = OxmlElement('w:tblPr')
                element._tbl.insert(0, tPr)
            tblInd = OxmlElement('w:tblInd')
            tblInd.set(qn('w:w'), '720')
            tblInd.set(qn('w:type'), 'dxa')
            tPr.append(tblInd)
        else:
            self._set_blockquote_indent(element)

    def _set_blockquote_indent(self, paragraph):
        """Set left indent and a left border on a paragraph for blockquote look."""
        pPr = paragraph._p.get_or_add_pPr()

        # Left indent (720 twips = 0.5 inch)
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '720')
        pPr.append(ind)

        # Grey left border
        pBdr = OxmlElement('w:pBdr')
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'single')
        left.set(qn('w:sz'), '12')
        left.set(qn('w:space'), '4')
        left.set(qn('w:color'), 'CCCCCC')
        pBdr.append(left)
        pPr.append(pBdr)

    def add_blockquote(self, bq_lines):
        """Process a block of consecutive blockquote lines.

        Strips the leading '>' from each line, detects tables and
        regular paragraphs inside, and renders them with blockquote styling.
        """
        # Strip '> ' or '>' prefix from each line
        stripped = []
        for line in bq_lines:
            if line.startswith('> '):
                stripped.append(line[2:])
            elif line.startswith('>'):
                stripped.append(line[1:])
            else:
                stripped.append(line)

        # Process the inner content
        i = 0
        while i < len(stripped):
            inner = stripped[i]
            inner_stripped = inner.strip()

            # Empty line — skip
            if not inner_stripped:
                i += 1
                continue

            # Table inside blockquote
            if inner_stripped.startswith('|'):
                table_lines = []
                while i < len(stripped) and stripped[i].strip().startswith('|'):
                    table_lines.append(stripped[i])
                    i += 1
                self.add_table(table_lines)
                # Apply blockquote indent to the table we just added
                self._apply_blockquote_style(self.doc.tables[-1])
                continue

            # Regular paragraph in blockquote
            p = self.add_paragraph(inner)
            self._apply_blockquote_style(p)
            i += 1

    def add_image(self, image_path, width=None):
        """Add an image to the document."""
        try:
            if width:
                self.doc.add_picture(image_path, width=Inches(width))
            else:
                self.doc.add_picture(image_path, width=Inches(5))
        except Exception as e:
            print(f"Warning: Could not add image {image_path}: {e}")

    def convert(self):
        """Convert markdown to DOCX"""
        print(f"Converting {self.input_file} to {self.output_file}...")

        # Read markdown content
        with open(self.input_file, 'r', encoding='utf-8') as f:
            md_content = f.read()

        # Process mermaid diagrams
        md_content = self.extract_and_convert_mermaid(md_content)

        # Parse markdown line by line for better control
        lines = md_content.split('\n')
        i = 0
        in_code_block = False
        code_lines = []
        code_lang = ''
        prev_was_numbered_list = False
        current_list_num_id = None

        while i < len(lines):
            line = lines[i]

            # Code blocks (support indented fences like '   ```json')
            stripped_line = line.strip()
            if stripped_line.startswith('```'):
                if not in_code_block:
                    in_code_block = True
                    code_lang = stripped_line[3:].strip()
                    code_lines = []
                else:
                    in_code_block = False
                    self.add_code_block('\n'.join(code_lines), code_lang)
                i += 1
                continue

            if in_code_block:
                code_lines.append(line)
                i += 1
                continue

            # Headings
            if line.startswith('#'):
                level = len(line) - len(line.lstrip('#'))
                text = line.lstrip('#').strip()
                if level <= 6:
                    self.add_heading(text, level)
                    i += 1
                    continue

            # Images
            img_match = re.match(r'!\[.*?\]\((.+?)\)', line)
            if img_match:
                img_path = img_match.group(1)
                # Handle relative paths
                if not os.path.isabs(img_path):
                    img_path = os.path.join(os.path.dirname(self.input_file), img_path)
                self.add_image(img_path)
                i += 1
                continue

            # Tables — collect consecutive pipe-delimited lines
            if line.strip().startswith('|'):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith('|'):
                    table_lines.append(lines[i])
                    i += 1
                self.add_table(table_lines)
                continue

            # Horizontal rule
            if line.strip() in ['---', '***', '___']:
                self.add_horizontal_rule()
                i += 1
                continue

            # Blockquotes — collect consecutive lines starting with '>'
            if stripped_line.startswith('>'):
                bq_lines = []
                while i < len(lines) and lines[i].strip().startswith('>'):
                    bq_lines.append(lines[i].strip())
                    i += 1
                self.add_blockquote(bq_lines)
                continue

            # Lists (simple implementation)
            if line.strip().startswith(('- ', '* ', '+ ')):
                text = line.strip()[2:]
                self.add_paragraph(text, style='List Bullet')
                i += 1
                continue

            # Numbered lists
            if re.match(r'^\d+\.\s', line.strip()):
                text = re.sub(r'^\d+\.\s', '', line.strip())
                p = self.add_paragraph(text, style='List Number')
                # Create a new numbering definition for the first item of each list
                if not prev_was_numbered_list:
                    current_list_num_id = self._create_numbering()
                # Apply the same numId to every item in this list group
                self._apply_numbering(p, current_list_num_id)
                prev_was_numbered_list = True
                i += 1
                continue

            prev_was_numbered_list = False
            current_list_num_id = None

            # Regular paragraph
            if line.strip():
                self.add_paragraph(line)
            else:
                # Empty line - could be paragraph break
                pass

            i += 1

        # Close the output file in Word if it's open (Windows only)
        self._close_in_word(self.output_file)

        # Save document
        self.doc.save(self.output_file)
        print(f"Successfully created {self.output_file}")

        # Cleanup temp files
        self.cleanup()

    @staticmethod
    def _close_in_word(docx_path):
        """If the docx is open in Word, close just that document (not all of Word).

        Uses Windows PowerShell 5.1 COM interop so it only works on Windows.
        Silently does nothing on other platforms or if Word isn't running.
        """
        if sys.platform != 'win32':
            return
        abs_path = str(Path(docx_path).resolve()).replace("'", "''")
        script = (
            "try {\n"
            "  $w = [System.Runtime.InteropServices.Marshal]::GetActiveObject('Word.Application')\n"
            f"  $doc = $w.Documents | Where-Object {{ $_.FullName -eq '{abs_path}' }}\n"
            "  if ($doc) {\n"
            "    $doc.Saved = $true\n"
            "    $doc.Close()\n"
            "    Write-Host 'Closed document in Word'\n"
            "    Start-Sleep -Milliseconds 500\n"
            "  }\n"
            "  [void][System.Runtime.InteropServices.Marshal]::ReleaseComObject($w)\n"
            "} catch { }\n"
        )
        # Write to a temp .ps1 file to avoid variable-escaping issues
        # between pwsh (PS 7) and powershell.exe (PS 5.1)
        script_file = tempfile.mktemp(suffix='.ps1')
        try:
            with open(script_file, 'w', encoding='utf-8') as f:
                f.write(script)
            subprocess.run(
                ['powershell.exe', '-NoProfile', '-ExecutionPolicy', 'Bypass', '-File', script_file],
                capture_output=True, timeout=15,
            )
        except (subprocess.SubprocessError, FileNotFoundError):
            pass
        finally:
            try:
                os.unlink(script_file)
            except OSError:
                pass


def main():
    parser = argparse.ArgumentParser(description='Convert Markdown to DOCX with Mermaid support')
    parser.add_argument('input', help='Input markdown file')
    parser.add_argument('output', nargs='?', help='Output DOCX file (default: input file with .docx extension)')

    args = parser.parse_args()

    if not os.path.exists(args.input):
        print(f"Error: Input file '{args.input}' not found")
        sys.exit(1)

    converter = MarkdownToDocxConverter(args.input, args.output)
    try:
        converter.convert()
    except Exception as e:
        print(f"Error during conversion: {e}")
        converter.cleanup()
        sys.exit(1)


if __name__ == '__main__':
    main()
